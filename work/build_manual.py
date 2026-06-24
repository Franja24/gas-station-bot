from pathlib import Path
import re

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/frankvargas/Documents/Gas Station petro 7")
SOURCE = Path(
    "/Users/frankvargas/Downloads/"
    "Manual de instalación y ejecución del proyecto Gas Station Automation.docx"
)
SCREENSHOT = ROOT / "evidence/happy-path/04-sevenly-benefits-prompt.png"
LOGO = ROOT / "work/sevenly_logo.png"
OUTPUT = ROOT / "output/Manual Gas Station Automation - Sevenly v1.0.0.docx"

GREEN = "008761"
DARK_GREEN = "005C46"
RED = "E21D2A"
INK = "202124"
MUTED = "66706B"
LIGHT_GREEN = "EAF5F1"
LIGHT_GRAY = "F3F5F4"
WHITE = "FFFFFF"

MANUAL_CONTENT = [
    "1. Objetivo y regla principal",
    "Cada persona debe calibrar el proyecto en una rama propia, sin trabajar directamente sobre main.",
    "La rama main debe mantenerse limpia, estable y actualizada con el repositorio remoto.",
    "Los cambios de coordenadas deben probarse en la misma resolución, escala, zoom y tamaño de ventana que se usarán normalmente.",
    "2. Requisitos previos",
    "Antes de comenzar, verificar que estén instalados:",
    "Git",
    "Python 3",
    "Terminal, iTerm o PyCharm",
    "La aplicación remota que se usará durante las pruebas, por ejemplo AnyDesk, RustDesk o Windows App",
    "Verificar versiones:",
    "git --version",
    "python3 --version",
    "3. Clonar el proyecto por primera vez",
    "Clonar el repositorio:",
    "git clone https://github.com/Franja24/gas-station-bot.git",
    "Entrar al proyecto:",
    "cd gas-station-bot",
    "Verificar el repositorio remoto:",
    "git remote -v",
    "Nota: este paso solo se realiza la primera vez en cada computadora.",
    "4. Mantener main actualizado antes de trabajar",
    "Antes de crear una rama nueva o comenzar una calibración, actualizar main:",
    "git switch main",
    "git pull --ff-only origin main",
    "Si git indica que existen cambios locales, no forzar la actualización. Guardar o revisar esos cambios antes de continuar.",
    "5. Crear una rama personal de calibración",
    "Crear una rama con el nombre de la persona y el equipo o pantalla:",
    "git switch -c feature/calibracion-frank-macbook",
    "Ejemplos recomendados:",
    "feature/calibracion-frank-macbook",
    "feature/calibracion-ana-monitor1080",
    "Verificar la rama actual:",
    "git branch --show-current",
    "No trabajar directamente sobre main.",
    "6. Continuar una rama existente y sincronizarla",
    "Si la rama ya existe localmente, cambiar a ella:",
    "git switch feature/calibracion-frank-macbook",
    "Traer los cambios recientes del servidor:",
    "git fetch origin",
    "Integrar main actualizado en la rama personal:",
    "git merge origin/main",
    "Si aparecen conflictos, resolverlos y volver a ejecutar las pruebas antes de subir cambios.",
    "7. Crear el entorno virtual e instalar dependencias",
    "Crear el entorno virtual:",
    "python3 -m venv venv",
    "Activar en Mac/Linux:",
    "source venv/bin/activate",
    "Activar en Windows:",
    "venv\\Scripts\\activate",
    "Instalar dependencias:",
    "pip install pyautogui opencv-python pillow mss numpy pynput behave reportlab",
    "Nota: la carpeta venv no debe subirse a Git.",
    "8. Habilitar permisos en macOS",
    "En macOS, habilitar permisos para la aplicación desde donde se ejecutará el proyecto:",
    "System Settings → Privacy & Security → Accessibility",
    "System Settings → Privacy & Security → Screen Recording",
    "Agregar la aplicación correspondiente:",
    "PyCharm",
    "Terminal",
    "iTerm",
    "Después de cambiar permisos, cerrar y volver a abrir la aplicación.",
    "9. Preparar la pantalla antes de calibrar",
    "Mantener constantes estas condiciones durante la calibración y la ejecución:",
    "Resolución del monitor",
    "Escalado de pantalla",
    "Tamaño y posición de la ventana",
    "AnyDesk / RustDesk / Windows App",
    "Zoom de la aplicación",
    "Nota: si cambia alguna de estas condiciones, las coordenadas pueden dejar de funcionar y será necesario recalibrar.",
    "10. Ejecutar una prueba inicial",
    "Ejecutar desde terminal:",
    "python main.py",
    "También puede ejecutarse desde PyCharm usando el botón Run.",
    "Antes de editar coordenadas, observar qué clic o búsqueda falla y registrar la posición esperada.",
    "11. Ajustar coordenadas según la pantalla",
    "Las coordenadas se ajustan actualmente en:",
    "clicker.py",
    "Buscar estructuras como:",
    "COORDS = {",
    "\"continue_button.png\": (800, 530),",
    "}",
    "Formato:",
    "(x, y)",
    "Regla rápida:",
    "Si el click cae muy a la izquierda → aumentar x",
    "Si el click cae muy a la derecha → disminuir x",
    "Si el click cae muy arriba → aumentar y",
    "Si el click cae muy abajo → disminuir y",
    "Modificar valores poco a poco, ejecutar nuevamente y conservar únicamente coordenadas verificadas.",
    "12. Ajustar regiones de búsqueda",
    "Las regiones de búsqueda también se ajustan actualmente en:",
    "clicker.py",
    "Ejemplo:",
    "REGIONS = {",
    "\"premium.png\": (900, 300, 600, 400)",
    "}",
    "Formato:",
    "(x_inicio, y_inicio, ancho, alto)",
    "Una región debe incluir completamente el elemento buscado sin abarcar una zona innecesariamente grande.",
    "Después de modificar una región, validar que la imagen se encuentre de forma consistente.",
    "13. Validar la calibración completa",
    "Ejecutar el flujo completo al menos dos veces con la misma configuración de pantalla.",
    "Si existen escenarios BDD, ejecutar:",
    "behave",
    "Para ejecutar solo un feature:",
    "behave features/premium.feature --no-capture",
    "Confirmar que los clics, búsquedas y transiciones funcionen sin intervención manual.",
    "14. Revisar evidencias y archivos antes de confirmar cambios",
    "Cada ejecución puede generar una carpeta dentro de:",
    "evidencias/",
    "También pueden generarse:",
    "screenshots/",
    "execution_report.pdf",
    "Revisar el estado del repositorio:",
    "git status",
    "Nota: no subir evidencias, venv, credenciales, direcciones IP ni archivos temporales a Git.",
    "15. Confirmar y subir únicamente los archivos necesarios",
    "Agregar de forma explícita solo los archivos que se modificaron intencionalmente:",
    "git add clicker.py",
    "Revisar exactamente lo que se incluirá en el commit:",
    "git diff --staged",
    "Crear el commit:",
    "git commit -m \"Calibrate coordinates for Frank MacBook\"",
    "Subir la rama personal:",
    "git push -u origin feature/calibracion-frank-macbook",
    "Evitar git add . cuando existan evidencias o archivos locales sin revisar.",
    "16. Crear Pull Request y mantener la rama",
    "Crear un Pull Request desde la rama personal hacia main.",
    "En la descripción, indicar la resolución, escala, aplicación remota y flujo validado.",
    "Antes de agregar cambios nuevos a una rama existente, volver a sincronizarla:",
    "git fetch origin",
    "git merge origin/main",
    "Resolver cualquier conflicto y repetir la validación completa.",
    "17. Recomendaciones importantes",
    "No trabajar directamente sobre main.",
    "No forzar pull, merge o push si existen conflictos que no se comprenden.",
    "No asumir que una calibración funcionará en otra pantalla o configuración.",
    "Conservar una rama por persona y contexto de calibración para facilitar revisión y mantenimiento.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GREEN, size="10"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)


def set_table_width(table, dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Arial", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_border(section, color=GREEN):
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgBorders"))
    if old is not None:
        sect_pr.remove(old)
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "16")
        border.set(qn("w:space"), "22")
        border.set(qn("w:color"), color)
        borders.append(border)
    sect_pr.append(borders)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def paragraph_border(paragraph, color=GREEN, size="10", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def create_logo():
    image = Image.open(SCREENSHOT).convert("RGBA")
    # The clean Sevenly wordmark shown above the order summary.
    logo = image.crop((510, 216, 710, 279))
    # Remove surrounding near-white pixels and crop to the visible wordmark.
    pixels = logo.load()
    for y in range(logo.height):
        for x in range(logo.width):
            r, g, b, a = pixels[x, y]
            if r > 238 and g > 238 and b > 238:
                pixels[x, y] = (255, 255, 255, 0)
    bbox = logo.getbbox()
    if bbox:
        logo = logo.crop(bbox)
    logo = logo.resize((logo.width * 4, logo.height * 4), Image.Resampling.LANCZOS)
    logo.save(LOGO)


def set_section_geometry(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.15)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)
    section.header_distance = Cm(0.85)
    section.footer_distance = Cm(0.85)
    add_page_border(section)


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    styles = doc.styles
    heading = styles["Heading 1"]
    heading.font.name = "Arial"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    heading.font.size = Pt(15)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Menlo"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Menlo")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Menlo")
    code.font.size = Pt(8.7)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = True

    label = styles.add_style("Instruction Label", WD_STYLE_TYPE.PARAGRAPH)
    label.font.name = "Arial"
    label._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    label._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    label.font.size = Pt(10)
    label.font.bold = True
    label.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(3)
    label.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Cm(0.65)
    bullet.paragraph_format.first_line_indent = Cm(-0.25)
    bullet.paragraph_format.space_after = Pt(3)
    bullet.paragraph_format.line_spacing = 1.1


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Cm(5.0))
    p.paragraph_format.space_after = Pt(22)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    set_run_font(kicker.add_run("MANUAL TÉCNICO"), size=10, color=RED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(
        title.add_run("Instalación, mantenimiento y calibración"),
        size=25,
        color=DARK_GREEN,
        bold=True,
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run_font(
        subtitle.add_run("Gas Station Automation"), size=17, color=INK, bold=True
    )

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_border(rule, color=RED, size="18", space="8")
    rule.paragraph_format.left_indent = Cm(4.0)
    rule.paragraph_format.right_indent = Cm(4.0)
    rule.paragraph_format.space_after = Pt(28)

    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    set_table_width(meta, 6200)
    meta.alignment = 1
    values = [("Versión", "1.0.0"), ("Elaborado por", "Frank Vargas")]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(7.0)
        for cell in row.cells:
            set_cell_shading(cell, LIGHT_GREEN)
            set_cell_border(cell, color=GREEN, size="8")
            cell.vertical_alignment = 1
        p1 = row.cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(p1.add_run(label), size=10, color=DARK_GREEN, bold=True)
        p2 = row.cells[1].paragraphs[0]
        set_run_font(p2.add_run(value), size=10.5, color=INK, bold=True)

    for _ in range(4):
        doc.add_paragraph()
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        footer_line.add_run("Guía para mantener el proyecto y calibrar cada pantalla"),
        size=9.5,
        color=MUTED,
        italic=True,
    )


def configure_running_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(LOGO), width=Cm(1.9))
    r = p.add_run("   Gas Station Automation | Manual técnico")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    paragraph_border(p, color=GREEN, size="7", space="3")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    set_run_font(p.add_run("Versión 1.0.0  |  Frank Vargas  |  Página "), size=8.5, color=MUTED)
    add_page_number(p)


def add_contents(doc, headings):
    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    title.add_run("Contenido")
    paragraph_border(title, color=GREEN, size="10", space="5")

    intro = doc.add_paragraph()
    intro.add_run(
        "Esta guía define un flujo seguro para mantener el proyecto actualizado, "
        "trabajar en ramas personales y calibrar coordenadas según cada pantalla."
    )

    for heading in headings:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(heading), size=10, color=DARK_GREEN, bold=True)
    doc.add_page_break()


CODE_PREFIXES = (
    "git ",
    "cd ",
    "python3 ",
    "python ",
    "source ",
    "venv\\",
    "pip ",
    "behave ",
    "COORDS =",
    "REGIONS =",
    '"',
    "}",
    "(",
    "* feature",
    "evidencias/",
    "screenshots/",
    "feature/",
    "System Settings ",
)

CODE_EXACT = {
    "behave",
    "main",
    "clicker.py",
    "execution_report.pdf",
    "Git",
    "Python 3",
}

BULLETS = {
    "Git",
    "Python 3",
    "Terminal, iTerm o PyCharm",
    "La aplicación remota que se usará durante las pruebas, por ejemplo AnyDesk, RustDesk o Windows App",
    "PyCharm",
    "Terminal",
    "iTerm",
    "Resolución del monitor",
    "Escalado de pantalla",
    "Tamaño de ventana",
    "AnyDesk / RustDesk",
    "Windows App",
    "AnyDesk / RustDesk / Windows App",
    "Zoom",
    "Si el click cae muy a la izquierda → aumentar x",
    "Si el click cae muy a la derecha → disminuir x",
    "Si el click cae muy arriba → aumentar y",
    "Si el click cae muy abajo → disminuir y",
}


def add_manual_content(doc, source_paragraphs):
    for text in source_paragraphs:
        if re.match(r"^\d+\.\s", text):
            if text.startswith(("12. ", "15. ")):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(text)
            paragraph_border(p, color=GREEN, size="8", space="4")
            continue

        if text in BULLETS:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
            continue

        if text in CODE_EXACT or text.startswith(CODE_PREFIXES):
            p = doc.add_paragraph(style="Code Block")
            p.add_run(text)
            shade_paragraph(p, LIGHT_GRAY)
            continue

        if (
            text.startswith("Nota:")
            or text.startswith("No trabajar directamente")
            or text.startswith("Evitar git add")
        ):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            set_table_width(table, 8200)
            cell = table.cell(0, 0)
            set_cell_shading(cell, LIGHT_GREEN)
            set_cell_border(cell, color=GREEN, size="10")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), size=10, color=DARK_GREEN, bold=True)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
            continue

        if text.endswith(":") or text in {"Ejemplo real:", "Regla rápida:"}:
            p = doc.add_paragraph(style="Instruction Label")
            p.add_run(text)
            continue

        p = doc.add_paragraph()
        p.add_run(text)


def main():
    ROOT.joinpath("work").mkdir(exist_ok=True)
    ROOT.joinpath("output").mkdir(exist_ok=True)
    create_logo()

    paragraphs = MANUAL_CONTENT
    source_title = "Manual de instalación, mantenimiento y calibración del proyecto Gas Station Automation"
    headings = [text for text in paragraphs if re.match(r"^\d+\.\s", text)]

    doc = Document()
    style_document(doc)
    first_section = doc.sections[0]
    set_section_geometry(first_section)
    first_section.different_first_page_header_footer = True
    add_cover(doc)

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(content_section)
    set_page_number_start(content_section, 1)
    configure_running_header_footer(content_section)

    add_contents(doc, headings)
    add_manual_content(doc, paragraphs)

    doc.core_properties.title = source_title
    doc.core_properties.subject = "Manual técnico de Gas Station Automation"
    doc.core_properties.author = "Frank Vargas"
    doc.core_properties.comments = "Versión 1.0.0"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
